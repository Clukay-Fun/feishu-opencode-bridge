#!/usr/bin/env python3
from pathlib import Path
import sys

from docx import Document  # type: ignore

from common.io import read_json_stdin, write_error_stderr, write_json_stdout


def extract_table_lines(document: Document) -> list[str]:
    lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        lines.append(text)
    return lines


def extract_docx_text(input_path: Path) -> str:
    document = Document(str(input_path))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    lines.extend(extract_table_lines(document))
    return "\n".join(lines)


def extract_text(input_path: Path) -> tuple[str, str]:
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(input_path), "docx"
    if suffix == ".txt":
        return input_path.read_text(encoding="utf-8"), "txt"
    if suffix == ".md":
        return input_path.read_text(encoding="utf-8"), "md"
    raise ValueError(f"unsupported file type: {suffix}")


def main() -> int:
    try:
        payload = read_json_stdin()
        input_path = Path(str(payload["inputPath"])).expanduser().resolve()
        if not input_path.exists():
            raise FileNotFoundError(f"input file not found: {input_path}")
        text, fmt = extract_text(input_path)
        write_json_stdout({
            "text": text,
            "format": fmt,
        })
        return 0
    except Exception as exc:  # noqa: BLE001
        write_error_stderr(str(exc))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
