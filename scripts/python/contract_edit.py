#!/usr/bin/env python3
"""
职责: 对合同 DOCX 执行小范围结构编辑，如删条款、替换内容等。
关注点:
- 用段落级操作实现工作台发起的轻量编辑。
- 返回已应用和跳过的操作，便于上游解释结果。
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from docx import Document  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

from common.io import read_json_stdin, write_error_stderr, write_json_stdout

CLAUSE_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十百零〇\d]+条"),
    re.compile(r"^第[一二三四五六七八九十百零〇\d]+款"),
    re.compile(r"^第[一二三四五六七八九十百零〇\d]+项"),
    re.compile(r"^附件[一二三四五六七八九十百零〇\d]*"),
    re.compile(r"^[一二三四五六七八九十]+、"),
    re.compile(r"^\(?[一二三四五六七八九十]+\)?[、.．]"),
    re.compile(r"^\(?\d+\)?[、.．]"),
    re.compile(r"^\d+(?:\.\d+)+"),
]


"""Wrap a raw XML paragraph element back into a python-docx `Paragraph`."""
def make_paragraph(element, parent) -> Paragraph:
    return Paragraph(element, parent)


"""Return the document's current paragraph list snapshot."""
def iter_paragraphs(document: Document) -> list[Paragraph]:
    return list(document.paragraphs)


"""Normalize paragraph text for matching and emptiness checks."""
def paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


"""Detect whether the paragraph looks like a clause heading."""
def is_clause_heading(text: str) -> bool:
    stripped = text.strip()
    return any(pattern.search(stripped) for pattern in CLAUSE_PATTERNS)


"""Return whether a paragraph ends a logical page in DOCX XML."""
def has_page_boundary(paragraph: Paragraph) -> bool:
    element = paragraph._element
    for br in element.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    for rendered_break in element.iter(qn("w:lastRenderedPageBreak")):
        if rendered_break is not None:
            return True
    for section in element.iter(qn("w:sectPr")):
        if section is not None:
            return True
    return False


"""Remove one paragraph from the document tree."""
def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


"""Insert a new paragraph after the anchor, optionally copying the body style."""
def insert_paragraph_after(anchor: Paragraph, text: str, style_name: Optional[str] = None) -> Paragraph:
    new_element = OxmlElement("w:p")
    anchor._element.addnext(new_element)
    paragraph = make_paragraph(new_element, anchor._parent)
    if style_name:
        paragraph.style = style_name
    paragraph.add_run(text)
    return paragraph


"""Find the paragraph range that belongs to a target clause."""
def find_clause_range(paragraphs: list[Paragraph], clause_number: Optional[str] = None, heading: Optional[str] = None):
    start = -1
    for index, paragraph in enumerate(paragraphs):
        text = paragraph_text(paragraph)
        if not text:
            continue
        matched = False
        if clause_number and text.startswith(clause_number):
            matched = True
        if heading and heading in text:
            matched = True
        if matched:
            start = index
            break
    if start < 0:
        return None

    end = len(paragraphs)
    for index in range(start + 1, len(paragraphs)):
        if is_clause_heading(paragraph_text(paragraphs[index])):
            end = index
            break
    return start, end


"""Delete the matching clause block from the document."""
def delete_clause(document: Document, clause_number: Optional[str], heading: Optional[str]) -> bool:
    paragraphs = iter_paragraphs(document)
    clause_range = find_clause_range(paragraphs, clause_number=clause_number, heading=heading)
    if not clause_range:
        return False
    start, end = clause_range
    for paragraph in paragraphs[start:end]:
        remove_paragraph(paragraph)
    return True


"""Replace the body paragraphs under a matching clause heading."""
def replace_clause_content(document: Document, clause_number: Optional[str], heading: Optional[str], new_content: str) -> bool:
    paragraphs = iter_paragraphs(document)
    clause_range = find_clause_range(paragraphs, clause_number=clause_number, heading=heading)
    if not clause_range:
        return False
    start, end = clause_range
    heading_paragraph = paragraphs[start]
    body_paragraphs = paragraphs[start + 1:end]
    style_name = body_paragraphs[0].style.name if body_paragraphs else None
    for paragraph in body_paragraphs:
        remove_paragraph(paragraph)

    anchor = heading_paragraph
    blocks = [line.strip() for line in str(new_content).replace("\r\n", "\n").split("\n") if line.strip()]
    if not blocks:
        return True
    for block in blocks:
        anchor = insert_paragraph_after(anchor, block, style_name)
    return True


"""Delete everything from the first matching heading to the end of the document."""
def delete_by_heading(document: Document, heading: str) -> bool:
    paragraphs = iter_paragraphs(document)
    start = -1
    for index, paragraph in enumerate(paragraphs):
        if heading in paragraph_text(paragraph):
            start = index
            break
    if start < 0:
        return False
    for paragraph in paragraphs[start:]:
        remove_paragraph(paragraph)
    return True


"""Delete logical pages split by explicit page/section breaks."""
def delete_pages(document: Document, page_range: list[int] | tuple[int, int]) -> bool:
    if len(page_range) != 2:
        return False
    start_page = int(page_range[0])
    end_page = int(page_range[1])
    if start_page <= 0 or end_page < start_page:
        return False

    paragraphs = iter_paragraphs(document)
    targets: list[Paragraph] = []
    current_page = 1
    for paragraph in paragraphs:
        if start_page <= current_page <= end_page:
            targets.append(paragraph)
        if has_page_boundary(paragraph):
            current_page += 1

    for paragraph in targets:
        remove_paragraph(paragraph)
    return len(targets) > 0


"""Apply supported edit operations and report which ones were skipped."""
def apply_operations(document: Document, operations: list[dict]) -> tuple[int, list[dict]]:
    applied = 0
    skipped: list[dict] = []
    for operation in operations:
        op_type = str(operation.get("type") or "").strip()
        ok = False
        if op_type == "delete_clause":
            ok = delete_clause(
                document,
                str(operation.get("clauseNumber")).strip() if operation.get("clauseNumber") else None,
                str(operation.get("heading")).strip() if operation.get("heading") else None,
            )
        elif op_type == "replace_content":
            new_content = str(operation.get("newContent") or "").strip()
            if new_content:
                ok = replace_clause_content(
                    document,
                    str(operation.get("clauseNumber")).strip() if operation.get("clauseNumber") else None,
                    str(operation.get("heading")).strip() if operation.get("heading") else None,
                    new_content,
                )
        elif op_type == "delete_by_heading":
            heading = str(operation.get("heading") or "").strip()
            if heading:
                ok = delete_by_heading(document, heading)
        elif op_type == "delete_pages":
            page_range = operation.get("pageRange")
            if isinstance(page_range, list) and len(page_range) == 2:
                ok = delete_pages(document, page_range)
        else:
            skipped.append({
                "operation": operation,
                "reason": f"unsupported operation: {op_type or 'unknown'}",
            })
            continue

        if ok:
            applied += 1
        else:
            skipped.append({
                "operation": operation,
                "reason": "target not found",
            })
    return applied, skipped


"""Read JSON input, edit the source DOCX, and emit edit results."""
def main() -> int:
    try:
        payload = read_json_stdin()
        input_path = Path(str(payload["inputPath"])).expanduser().resolve()
        output_path = Path(str(payload.get("outputPath") or input_path.with_name(f"{input_path.stem}-edited{input_path.suffix}"))).expanduser().resolve()
        operations = payload.get("operations")
        if not input_path.exists():
            raise FileNotFoundError(f"input file not found: {input_path}")
        if input_path.suffix.lower() != ".docx":
            raise ValueError("contract_edit 目前仅支持 .docx 文件")
        if not isinstance(operations, list) or not operations:
            raise ValueError("operations 不能为空")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document(str(input_path))
        applied_ops, skipped_ops = apply_operations(
            document,
            [item for item in operations if isinstance(item, dict)],
        )
        document.save(str(output_path))
        write_json_stdout({
            "outputPath": str(output_path),
            "appliedOps": applied_ops,
            "skippedOps": skipped_ops,
        })
        return 0
    except Exception as exc:  # noqa: BLE001
        write_error_stderr(str(exc))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
