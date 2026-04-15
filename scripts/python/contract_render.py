#!/usr/bin/env python3
from pathlib import Path

from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.shared import Pt  # type: ignore

from common.io import read_json_stdin, write_error_stderr, write_json_stdout
from common.styles import (
    BODY_FIRST_LINE_INDENT_PT,
    BODY_FONT_NAME,
    BODY_FONT_SIZE_PT,
    HEADING_FONT_NAME,
    HEADING_FONT_SIZE_PT,
    TITLE_FONT_NAME,
    TITLE_FONT_SIZE_PT,
)


def apply_default_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT_NAME
    normal.font.size = Pt(BODY_FONT_SIZE_PT)


def add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = TITLE_FONT_NAME
    run.font.size = Pt(TITLE_FONT_SIZE_PT)


def add_meta(document: Document, parties: dict) -> None:
    lines = [
        f"甲方：{parties.get('clientName') or '【待补】'}",
        f"对方：{parties.get('counterpartyName') or '【待补】'}",
    ]
    if parties.get("agencyName"):
        lines.append(f"乙方/机构：{parties['agencyName']}")
    if parties.get("leadLawyer"):
        lines.append(f"承办律师：{parties['leadLawyer']}")
    if parties.get("signDate"):
        lines.append(f"签署日期：{parties['signDate']}")
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.add_run(line)


def add_clause(document: Document, clause: dict) -> None:
    title = f"{clause.get('number') or ''} {clause.get('title') or ''}".strip()
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_run = heading.add_run(title or "未命名条款")
    heading_run.bold = True
    heading_run.font.name = HEADING_FONT_NAME
    heading_run.font.size = Pt(HEADING_FONT_SIZE_PT)

    content = str(clause.get("content") or "【待补】").replace("\r\n", "\n")
    for block in [segment.strip() for segment in content.split("\n") if segment.strip()]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(BODY_FIRST_LINE_INDENT_PT)
        paragraph.add_run(block)


def add_appendices(document: Document, appendices: list) -> None:
    if not appendices:
        return
    heading = document.add_paragraph()
    run = heading.add_run("附件")
    run.bold = True
    run.font.name = HEADING_FONT_NAME
    run.font.size = Pt(HEADING_FONT_SIZE_PT)
    for appendix in appendices:
        title = appendix.get("title") or "未命名附件"
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.name = HEADING_FONT_NAME
        title_run.font.size = Pt(HEADING_FONT_SIZE_PT)
        content = str(appendix.get("content") or "").replace("\r\n", "\n")
        for block in [segment.strip() for segment in content.split("\n") if segment.strip()]:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(BODY_FIRST_LINE_INDENT_PT)
            paragraph.add_run(block)


def render_document(state: dict, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    apply_default_style(document)
    add_title(document, str(state.get("title") or "合同草稿"))
    document.add_paragraph("")
    add_meta(document, state.get("parties") or {})
    document.add_paragraph("")

    clauses = state.get("clauses") or []
    for clause in clauses:
        if isinstance(clause, dict):
            add_clause(document, clause)

    appendices = state.get("appendices") or []
    filtered = [item for item in appendices if isinstance(item, dict)]
    if filtered:
        document.add_page_break()
        add_appendices(document, filtered)

    document.save(str(output_path))


def main() -> int:
    try:
        payload = read_json_stdin()
        state = payload.get("state") or {}
        output_path = Path(str(payload["outputPath"])).expanduser().resolve()
        render_document(state, output_path)
        write_json_stdout({
            "outputPath": str(output_path),
            "pageCount": None,
        })
        return 0
    except Exception as exc:  # noqa: BLE001
        write_error_stderr(str(exc))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
