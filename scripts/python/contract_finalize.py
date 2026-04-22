#!/usr/bin/env python3
"""
职责: 基于模板合同和字段数据生成接近交付态的最终 DOCX。
关注点:
- 直接改写模板中的关键占位段落。
- 清理签署区、收费方式和风险代理告知书等可选内容。
"""
from copy import deepcopy
from pathlib import Path

from docx import Document  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.table import Table  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

from common.io import read_json_stdin, write_error_stderr, write_json_stdout


"""Return trimmed paragraph text."""
def get_text(paragraph):
    return paragraph.text.strip()


"""Normalize text for fuzzy matching across spaces and line breaks."""
def normalize_match_text(text: str) -> str:
    return "".join(text.split())


"""Replace a paragraph's runs with new text while preserving the paragraph node."""
def clear_and_set_text(paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        run.clear()
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


"""Delete one paragraph from the XML tree."""
def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


"""Insert a new paragraph after an anchor and clone paragraph properties when possible."""
def insert_paragraph_after(anchor, text: str, style_source=None):
    new_element = OxmlElement("w:p")
    source = style_source if style_source is not None else anchor
    if source._element.pPr is not None:
        new_element.append(deepcopy(source._element.pPr))
    anchor._element.addnext(new_element)
    paragraph = Paragraph(new_element, anchor._parent)
    paragraph.add_run(text)
    return paragraph


"""Yield paragraphs from a document-like container, including nested tables."""
def iter_container_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in getattr(container, "tables", []):
        yield from iter_table_paragraphs(table)


"""Yield all paragraphs nested inside a table."""
def iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            yield from iter_container_paragraphs(cell)


"""Detect paragraphs around the signature block that need empty-line cleanup."""
def is_signature_cleanup_anchor(text: str) -> bool:
    normalized = normalize_match_text(text)
    anchors = [
        "甲方：",
        "乙方：北京市隆安（深圳）律师事务所",
        "法定代表人/负责人/授权代表：",
        "承办律师：",
        "签约时间：",
        "委托人：",
        "（以下无正文，为本合同签署处",
        "委托代理合同",
        "合同编号：",
        "附：《风险代理告知书》",
        "附件",
    ]
    return any(
        text.startswith(anchor)
        or anchor in text
        or normalize_match_text(anchor) in normalized
        for anchor in anchors
    )


"""Remove empty paragraphs that only pad the signature or appendix area."""
def remove_targeted_empty_paragraphs(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    for index in range(len(paragraphs) - 1, -1, -1):
        paragraph = paragraphs[index]
        if get_text(paragraph):
            continue
        prev_text = ""
        next_text = ""
        for prev_index in range(index - 1, -1, -1):
            prev_text = get_text(paragraphs[prev_index])
            if prev_text:
                break
        for next_index in range(index + 1, len(paragraphs)):
            next_text = get_text(paragraphs[next_index])
            if next_text:
                break
        if is_signature_cleanup_anchor(prev_text) or is_signature_cleanup_anchor(next_text):
            delete_paragraph(paragraph)


"""Append a signature block when the template does not already contain one."""
def ensure_signature_block(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    has_signature_placeholders = any(
        "法定代表人/负责人/授权代表：__________" in get_text(paragraph)
        or "承办律师：_____________" in get_text(paragraph)
        for paragraph in paragraphs
    )
    if has_signature_placeholders:
        return

    signature_anchor = None
    for paragraph in paragraphs:
        text = get_text(paragraph)
        if text.startswith("2.甲方在本合同签署之前已经全部阅读并知悉合同内容"):
            signature_anchor = paragraph
    if signature_anchor is None:
        signature_anchor = next((paragraph for paragraph in paragraphs if get_text(paragraph).startswith("第十二条 特别约定")), None)
    if signature_anchor is None:
        return

    anchor = insert_paragraph_after(signature_anchor, "甲方：", signature_anchor)
    anchor = insert_paragraph_after(anchor, "法定代表人/负责人/授权代表：__________", signature_anchor)
    anchor = insert_paragraph_after(anchor, "", signature_anchor)
    anchor = insert_paragraph_after(anchor, "乙方：北京市隆安（深圳）律师事务所", signature_anchor)
    anchor = insert_paragraph_after(anchor, "承办律师：_____________", signature_anchor)
    insert_paragraph_after(anchor, "签约时间：  202 年   月     日", signature_anchor)


"""Rewrite the template document in place based on the finalized contract data."""
def process_document(document: Document, data: dict) -> None:
    client_name = str(data.get("client_name") or "").strip()
    client_id_code = str(data.get("client_id_code") or "").strip()
    client_address = str(data.get("client_address") or "").strip()
    client_phone = str(data.get("client_phone") or "").strip()
    client_representative = str(data.get("client_representative") or "").strip()
    lead_lawyer = str(data.get("lead_lawyer") or "").strip()
    counterparty_name = str(data.get("counterparty_name") or "").strip()
    case_cause = str(data.get("case_cause") or "").strip()

    is_company = bool(data.get("is_company"))
    show_risk_notice = bool(data.get("show_risk_notice"))

    keep_stage = {
        "仲裁阶段": bool(data.get("engage_arbitration")),
        "一审诉讼": bool(data.get("engage_first_instance")),
        "二审诉讼": bool(data.get("engage_second_instance")),
        "执行程序": bool(data.get("engage_enforcement")),
        "调解、和解事宜": bool(data.get("engage_settlement")),
    }

    stage_fixed = bool(data.get("is_stage_fixed"))
    stage_fee_clauses = {
        "1.仲裁阶段：": str(data.get("fee_arbitration_clause") or "").strip(),
        "2.一审阶段：": str(data.get("fee_first_instance_clause") or "").strip(),
        "3.二审阶段：": str(data.get("fee_second_instance_clause") or "").strip(),
        "4.执行阶段：": str(data.get("fee_enforcement_clause") or "").strip(),
    }

    risk_blocks = [
        "☐基础收费+风险收费",
        "☑基础收费+风险收费",
        "甲乙双方选择风险代理收费方式",
        "1.基础费用：",
        "2.风险收费：",
        "如果甲方实际收到的是现金以外的有形资产或财产权益",
        "对于风险收费，甲方只要有回款或有回收其他有形资产或财产权益",
        "本代理为风险收费方式",
    ]

    paragraphs = list(iter_container_paragraphs(document))
    delete_from_index = None

    for index, paragraph in enumerate(paragraphs):
        text = get_text(paragraph)
        if not text:
            continue

        if text.startswith("聘请方（甲方）："):
            clear_and_set_text(paragraph, f"聘请方（甲方）：{client_name}")
            continue
        if "乙方：北京市隆安（深圳）律师事务所" in text and text.startswith("甲方："):
            clear_and_set_text(paragraph, "甲方：                                   乙方：北京市隆安（深圳）律师事务所")
            continue
        if text == "甲方：" or text.startswith("甲方："):
            clear_and_set_text(paragraph, f"甲方：{client_name}")
            continue
        if text.startswith("法定代表人/负责人："):
            if is_company:
                clear_and_set_text(paragraph, f"法定代表人/负责人：{client_representative}")
            else:
                delete_paragraph(paragraph)
            continue
        if text.startswith("证件号/社会统一信用代码："):
            clear_and_set_text(paragraph, f"证件号/社会统一信用代码：{client_id_code}")
            continue
        if text.startswith("地址："):
            clear_and_set_text(paragraph, f"地址：{client_address}")
            continue
        if text.startswith("联系电话："):
            clear_and_set_text(paragraph, f"联系电话：{client_phone}")
            continue
        if text.startswith("甲方因与") and "纠纷案件" in text:
            clear_and_set_text(paragraph, f"甲方因与{counterparty_name}{case_cause}纠纷案件，委托乙方代理，经双方协商，订立下列各条款，共同遵照履行。")
            continue
        if text.startswith("（一）乙方指派"):
            clear_and_set_text(paragraph, f"（一）乙方指派{lead_lawyer}律师作为案件中甲方的委托代理人，甲方同意上述律师指派其他律师和助理配合完成辅助工作，但乙方更换代理律师应取得甲方认可。")
            continue
        if text.endswith("仲裁阶段；"):
            if keep_stage["仲裁阶段"]:
                clear_and_set_text(paragraph, "☑仲裁阶段；")
            else:
                delete_paragraph(paragraph)
            continue
        if text.endswith("一审诉讼；"):
            if keep_stage["一审诉讼"]:
                clear_and_set_text(paragraph, "☑一审诉讼；")
            else:
                delete_paragraph(paragraph)
            continue
        if text.endswith("二审诉讼；"):
            if keep_stage["二审诉讼"]:
                clear_and_set_text(paragraph, "☑二审诉讼；")
            else:
                delete_paragraph(paragraph)
            continue
        if text.endswith("执行程序；"):
            if keep_stage["执行程序"]:
                clear_and_set_text(paragraph, "☑执行程序；")
            else:
                delete_paragraph(paragraph)
            continue
        if "调解、和解事宜" in text:
            if keep_stage["调解、和解事宜"]:
                clear_and_set_text(paragraph, "☑上述案件代理程序中，有关调解、和解事宜。")
            else:
                delete_paragraph(paragraph)
            continue
        if text == "☑按阶段收费" or text == "□按阶段收费":
            if stage_fixed:
                clear_and_set_text(paragraph, "☑按阶段收费")
            else:
                delete_paragraph(paragraph)
            continue
        if text.startswith("甲乙双方约定乙方律师费如下："):
            if not stage_fixed:
                delete_paragraph(paragraph)
            continue
        for prefix, replacement in stage_fee_clauses.items():
            if text.startswith(prefix):
                if stage_fixed and replacement:
                    clear_and_set_text(paragraph, replacement)
                else:
                    delete_paragraph(paragraph)
                break
        else:
            if stage_fixed and any(text.startswith(prefix) for prefix in risk_blocks):
                delete_paragraph(paragraph)
                continue
            if (not stage_fixed) and (text == "☑按阶段收费" or text.startswith("甲乙双方约定乙方律师费如下：")):
                delete_paragraph(paragraph)
                continue
            if "法定代表人/负责人/授权代表：" in normalize_match_text(text):
                clear_and_set_text(paragraph, "法定代表人/负责人/授权代表：__________   承办律师：_____________")
                continue
            if text.startswith("签约时间："):
                clear_and_set_text(paragraph, "签约时间：  202 年   月     日           ")
                continue
            if text.startswith("委托人："):
                clear_and_set_text(paragraph, "委托人：")
                continue
            if text.startswith("附：《风险代理告知书》") and not show_risk_notice:
                if delete_from_index is None:
                    delete_from_index = index
                continue
            if "风险代理告知书" in text and not show_risk_notice:
                if delete_from_index is None:
                    delete_from_index = index - 1 if index > 0 else index
                continue

    if not show_risk_notice:
        remove_tail = False
        for paragraph in list(document.paragraphs):
            text = get_text(paragraph)
            if (
                text.startswith("附：《风险代理告知书》")
                or text == "附件"
                or "风险代理告知书" in text
            ):
                remove_tail = True
            if remove_tail:
                delete_paragraph(paragraph)
    elif delete_from_index is not None:
        current = list(document.paragraphs)
        start = max(0, delete_from_index)
        for paragraph in current[start:]:
            delete_paragraph(paragraph)

    ensure_signature_block(document)
    remove_targeted_empty_paragraphs(document)


"""Read JSON input, finalize the contract document, and emit the saved path."""
def main() -> int:
    try:
        payload = read_json_stdin()
        input_path = Path(str(payload["inputPath"])).expanduser().resolve()
        output_path = Path(str(payload.get("outputPath") or input_path)).expanduser().resolve()
        data = payload.get("data")
        if not input_path.exists():
            raise FileNotFoundError(f"input file not found: {input_path}")
        if input_path.suffix.lower() != ".docx":
            raise ValueError("contract_finalize 目前仅支持 .docx 文件")
        if not isinstance(data, dict):
            raise ValueError("data 不能为空")
        document = Document(str(input_path))
        process_document(document, data)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        write_json_stdout({"outputPath": str(output_path)})
        return 0
    except Exception as exc:  # noqa: BLE001
        write_error_stderr(str(exc))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
